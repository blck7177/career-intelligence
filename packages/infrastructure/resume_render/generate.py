"""
Render a ResumeDocument into a formatted .docx file using python-docx.

Bullets use the built-in "List Bullet" paragraph style (backed by proper
w:numPr numbering defined in the style, not a literal "•" character) —
never insert bullet glyphs directly into run text.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .layout import DEFAULT_LAYOUT, LayoutConfig
from .schema import ResumeDocument, ResumeSection

_HEADING_RULE_COLOR = "444444"
_SUBTITLE_COLOR = "444444"
_BOLD_LEADIN_RE = re.compile(r"^\*\*(.+?)\*\*(.*)$", re.DOTALL)


def generate_docx(doc_content: ResumeDocument, output_path: str | Path, layout: LayoutConfig = DEFAULT_LAYOUT) -> Path:
    output_path = Path(output_path)
    document = Document()
    _apply_page_setup(document, layout)
    _apply_default_font(document, layout)

    _add_header_block(document, doc_content, layout)

    for section in doc_content.sections:
        _add_section(document, section, layout)

    document.save(output_path)
    return output_path


def _apply_page_setup(document: Document, layout: LayoutConfig) -> None:
    section = document.sections[0]
    section.top_margin = Inches(layout.margin_top_in)
    section.bottom_margin = Inches(layout.margin_bottom_in)
    section.left_margin = Inches(layout.margin_left_in)
    section.right_margin = Inches(layout.margin_right_in)


def _apply_default_font(document: Document, layout: LayoutConfig) -> None:
    normal = document.styles["Normal"]
    normal.font.name = layout.font_name
    normal.font.size = Pt(layout.body_size_pt)
    # East-Asian font fallback must be set via rFonts/eastAsia explicitly,
    # otherwise python-docx leaves it pointing at the template's default.
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), layout.font_name)


def _content_width_in(layout: LayoutConfig) -> float:
    return 8.5 - layout.margin_left_in - layout.margin_right_in


def _add_bottom_border(paragraph, color: str = _HEADING_RULE_COLOR, size: int = 6) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_header_block(document: Document, doc_content: ResumeDocument, layout: LayoutConfig) -> None:
    name_p = document.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    run = name_p.add_run(doc_content.name)
    run.bold = True
    run.font.size = Pt(layout.name_size_pt)

    contact_p = document.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(2)
    contact_p.add_run(doc_content.contact_line)

    if doc_content.headline:
        headline_p = document.add_paragraph()
        headline_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        headline_p.paragraph_format.space_after = Pt(12)
        run = headline_p.add_run(doc_content.headline)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(_SUBTITLE_COLOR.upper())


def _add_section(document: Document, section: ResumeSection, layout: LayoutConfig) -> None:
    heading_p = document.add_paragraph()
    # w:pBdr must precede w:spacing in <w:pPr> per the OOXML schema — add the
    # border before setting spacing, not after (python-docx appends each new
    # pPr child at the end, so call order here determines XML order).
    _add_bottom_border(heading_p)
    heading_p.paragraph_format.space_before = Pt(12)
    heading_p.paragraph_format.space_after = Pt(6)
    run = heading_p.add_run(section.heading)
    run.bold = True
    run.font.size = Pt(layout.heading_size_pt)

    for para_text in section.paragraphs:
        _add_body_paragraph(document, para_text)

    for entry in section.entries:
        _add_entry(document, entry, layout)

    for bullet_text in section.bullets:
        _add_bullet(document, bullet_text)


def _add_body_paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    match = _BOLD_LEADIN_RE.match(text)
    if match:
        bold_part, rest = match.groups()
        bold_run = p.add_run(bold_part)
        bold_run.bold = True
        if rest:
            p.add_run(rest)
    else:
        p.add_run(text)


def _add_entry(document: Document, entry, layout: LayoutConfig) -> None:
    title_p = document.add_paragraph()
    title_p.paragraph_format.space_before = Pt(8)
    title_p.paragraph_format.space_after = Pt(2)
    title_p.paragraph_format.tab_stops.add_tab_stop(
        Inches(_content_width_in(layout)), WD_TAB_ALIGNMENT.RIGHT
    )
    title_run = title_p.add_run(entry.title_line)
    title_run.bold = True
    if entry.dates:
        dates_run = title_p.add_run(f"\t{entry.dates}")
        dates_run.italic = True

    if entry.subtitle_line:
        subtitle_p = document.add_paragraph()
        subtitle_p.paragraph_format.space_after = Pt(4)
        subtitle_run = subtitle_p.add_run(entry.subtitle_line)
        subtitle_run.italic = True
        subtitle_run.font.color.rgb = RGBColor.from_string(_SUBTITLE_COLOR.upper())

    for bullet_text in entry.bullets:
        _add_bullet(document, bullet_text)


def _add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(6)
    p.add_run(text)
